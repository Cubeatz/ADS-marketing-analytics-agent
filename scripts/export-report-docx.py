#!/usr/bin/env python3
"""将 reports/{date}/daily-report.md 导出为 DOCX。"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

try:
    from workspace_lib import (
        document_dir,
        delivery_mode,
        feishu_webhook_configured,
        load_workspace,
        report_md_dir,
        sync_legacy_feishu_json,
    )
except ImportError:
    import sys
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from workspace_lib import (
        document_dir,
        delivery_mode,
        feishu_webhook_configured,
        load_workspace,
        report_md_dir,
        sync_legacy_feishu_json,
    )


def load_feishu_config(project_root: Path) -> dict:
    path = project_root / "config" / "feishu.json"
    if not path.exists():
        path = project_root / "config" / "feishu.example.json"
    if not path.exists():
        return {}
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def resolve_output_path(project_root: Path, date: str, feishu_cfg: dict) -> Path:
    workspace = load_workspace(project_root)
    if workspace:
        sync_legacy_feishu_json(project_root, workspace)
        delivery = workspace.get("delivery") or {}
        filename = delivery.get("local_docx_filename") or "daily-report.docx"
        out_dir = document_dir(project_root, workspace, date)
        out_dir.mkdir(parents=True, exist_ok=True)
        return out_dir / filename

    fallback = feishu_cfg.get("docx_fallback") or {}
    if not fallback.get("enabled", True):
        raise RuntimeError("docx_fallback.enabled 为 false，已跳过 DOCX 生成")

    filename = fallback.get("filename") or "daily-report.docx"
    custom = (fallback.get("custom_output_dir") or "").strip()
    if custom:
        out_dir = Path(custom)
    else:
        rel = (fallback.get("output_dir") or "reports/{date}").replace("{date}", date)
        out_dir = project_root / rel

    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir / filename


def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|")
    if not s:
        return False
    return bool(re.fullmatch(r"[\s\-\|:]+", s))


def parse_table_row(line: str) -> list[str]:
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def markdown_to_docx(md_text: str, out_path: Path, title_prefix: str, date: str) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "缺少 python-docx，请运行: pip install -r requirements.txt"
        ) from e

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    title = doc.add_heading(f"{title_prefix} {date}", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                if not is_table_separator(row_line):
                    rows.append(parse_table_row(row_line))
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        table.rows[r_idx].cells[c_idx].text = cell_text
                doc.add_paragraph("")
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph(stripped.lstrip("> ").strip())
            p.style = "Intense Quote"
            i += 1
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    doc.save(str(out_path))


def feishu_is_configured(feishu_cfg: dict) -> bool:
    if not feishu_cfg.get("enabled", False):
        return False
    url = (feishu_cfg.get("webhook_url") or "").strip()
    if not url:
        return False
    placeholders = ("YOUR_WEBHOOK", "YOUR_WEBHOOK_TOKEN", "example.com", "placeholder")
    return not any(p.lower() in url.lower() for p in placeholders)


def main() -> int:
    parser = argparse.ArgumentParser(description="导出营销日报为 DOCX")
    parser.add_argument("--date", required=True, help="报告日期 YYYY-MM-DD")
    parser.add_argument(
        "--project-root",
        default=None,
        help="项目根目录，默认自动推断",
    )
    parser.add_argument(
        "--markdown",
        default=None,
        help="Markdown 源文件，默认 reports/{date}/daily-report.md",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="输出 DOCX 路径，默认读 config/feishu.json docx_fallback",
    )
    args = parser.parse_args()

    script_dir = Path(__file__).resolve().parent
    project_root = Path(args.project_root) if args.project_root else script_dir.parent

    md_path = Path(args.markdown) if args.markdown else report_md_dir(project_root, load_workspace(project_root), args.date) / "daily-report.md"
    if not md_path.exists():
        md_path = project_root / "reports" / args.date / "daily-report.md"
    if not md_path.exists():
        print(f"错误: 找不到 Markdown 报告 {md_path}", file=sys.stderr)
        return 1

    feishu_cfg = load_feishu_config(project_root)
    out_path = Path(args.output) if args.output else resolve_output_path(project_root, args.date, feishu_cfg)

    md_text = md_path.read_text(encoding="utf-8")
    prefix = feishu_cfg.get("report_title_prefix") or "【营销日报】"
    markdown_to_docx(md_text, out_path, prefix, args.date)

    print(str(out_path.resolve()))
    return 0


if __name__ == "__main__":
    sys.exit(main())
